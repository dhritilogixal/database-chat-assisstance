#!/usr/bin/env python
# coding: utf-8

# ══════════════════════════════════════════════════════════════
# SUPERSTORE SALES DATABASE CHATBOT — Refactored
# Changes from previous version:
#   1. .env file support for all config (URL, model, DB)
#   2. Single call_model() utility function
#   3. Intent router — classifies every question
#   4. Mistral for generic/contextual/recommendation
#   5. Gemma3 only for SQL generation
#   6. Customers table JOIN support
#   7. Recommendation + subjective question handling
# ══════════════════════════════════════════════════════════════


import requests
import pandas as pd
import re
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from sqlalchemy import create_engine, text
from datetime import datetime, timedelta
import textwrap
import os
import io

# ── LOAD .env FILE ────────────────────────────────────────────
# Create a .env file on your Desktop with these values:
#   OLLAMA_URL=http://192.168.15.220:11434
#   SQL_MODEL=gemma3:12b
#   FAST_MODEL=mistral
#   DB_URL=postgresql://postgres:11223344@localhost:5432/sales_db

def load_env():
    env = {}
    env_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '.env')
    if os.path.exists(env_path):
        with open(env_path) as f:
            for line in f:
                line = line.strip()
                if line and not line.startswith('#') and '=' in line:
                    key, val = line.split('=', 1)
                    env[key.strip()] = val.strip()
        print("  ✅ .env file loaded")
    else:
        print("  ⚠️  No .env file found — using defaults")
    return env

env = load_env()

# ── CONFIG ────────────────────────────────────────────────────
OLLAMA_URL  = env.get('OLLAMA_URL', 'http://192.168.15.220:11434')
SQL_MODEL   = env.get('SQL_MODEL',  'gemma3:12b')   # for SQL generation
FAST_MODEL  = env.get('FAST_MODEL', 'mistral')      # for chat/recommendation
DB_URL      = env.get('DB_URL',     'postgresql://postgres:11223344@localhost:5432/sales_db')

# ── Word doc export ───────────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ── DATABASE ──────────────────────────────────────────────────
engine = create_engine(DB_URL)

print()
print("=" * 60)
print("       SUPERSTORE SALES DATABASE CHATBOT")
print("=" * 60)
print()

try:
    with engine.connect() as conn:
        conn.execute(text("SELECT 1"))
    print("  ✅ Database connected successfully")
except Exception as e:
    print(f"  ❌ Database connection failed: {e}")

try:
    requests.post(
        url=f"{OLLAMA_URL}/api/generate",
        headers={"Content-Type": "application/json"},
        json={"model": SQL_MODEL, "prompt": "hi", "stream": False,
              "options": {"temperature": 0.1, "num_ctx": 512}},
        timeout=60
    )
    print(f"  ✅ AI model connected ({SQL_MODEL})")
except Exception as e:
    print(f"  ❌ AI model connection failed: {e}")

if DOCX_AVAILABLE:
    print("  ✅ Word export available")
else:
    print("  ⚠️  Word export unavailable — pip install python-docx")

print()
print("=" * 60)

context = []
latest_chart_buf = None


# ══════════════════════════════════════════════════════════════
# UTILITY FUNCTION — call_model()
# Single function to call any Ollama model
# All model calls go through here — change URL/model in .env
# ══════════════════════════════════════════════════════════════

def call_model(prompt, model=None, temperature=0.4, top_p=0.9,
               repeat_penalty=1.1, num_ctx=4096, timeout=120):
    """
    Utility function to call any Ollama model.
    Uses FAST_MODEL by default (mistral).
    Pass model=SQL_MODEL for SQL generation.
    """
    if model is None:
        model = FAST_MODEL

    try:
        response = requests.post(
            url=f"{OLLAMA_URL}/api/generate",
            headers={"Content-Type": "application/json"},
            json={
                "model": model,
                "prompt": prompt,
                "stream": False,
                "options": {
                    "temperature": temperature,
                    "top_p": top_p,
                    "repeat_penalty": repeat_penalty,
                    "num_ctx": num_ctx
                }
            },
            timeout=timeout
        )
        return response.json()['response'].strip()
    except Exception as e:
        print(f"  ❌ Model call failed ({model}): {e}")
        return ""


# ══════════════════════════════════════════════════════════════
# DATABASE METADATA
# ══════════════════════════════════════════════════════════════

def get_metadata():
    metadata = "DATABASE STRUCTURE:\n"
    try:
        with engine.connect() as conn:
            tables = conn.execute(text("""
                SELECT table_name FROM information_schema.tables
                WHERE table_schema = 'public' ORDER BY table_name
            """)).fetchall()
            for table in tables:
                table_name = table[0]
                metadata += f"\nTable: {table_name}\n"
                metadata += "Columns:\n"
                columns = conn.execute(text(f"""
                    SELECT column_name, data_type
                    FROM information_schema.columns
                    WHERE table_name = '{table_name}'
                    ORDER BY ordinal_position
                """)).fetchall()
                for col in columns:
                    metadata += f"  - {col[0]} ({col[1]})\n"
                count = conn.execute(text(f"SELECT COUNT(*) FROM {table_name}")).fetchone()[0]
                metadata += f"Total rows: {count}\n"
    except Exception as e:
        print(f"  ❌ Metadata error: {e}")
    return metadata

print("\n  📋 Database metadata loaded:")
metadata = get_metadata()
print(metadata)
print("=" * 60)

def get_special_columns():
    special_cols = {}
    try:
        with engine.connect() as conn:
            tables = conn.execute(text("""
                SELECT table_name, column_name FROM information_schema.columns
                WHERE table_schema = 'public' AND column_name ~ '[^a-zA-Z0-9_]'
            """)).fetchall()
            for table_name, col_name in tables:
                normalized = col_name.replace('-', '_').replace(' ', '_').replace('.', '_')
                special_cols[normalized] = f'"{col_name}"'
                special_cols[col_name] = f'"{col_name}"'
    except Exception as e:
        print(f"  ⚠️ Could not get special columns: {e}")
    return special_cols

special_columns = get_special_columns()
print(f"  ✅ Special columns detected: {special_columns}")

def get_table_columns():
    table_cols = {}
    try:
        with engine.connect() as conn:
            tables = conn.execute(text("""
                SELECT table_name, column_name FROM information_schema.columns
                WHERE table_schema = 'public' ORDER BY table_name, ordinal_position
            """)).fetchall()
            for table_name, col_name in tables:
                if table_name not in table_cols:
                    table_cols[table_name] = []
                table_cols[table_name].append(col_name.lower())
    except Exception as e:
        print(f"  ⚠️ Could not get table columns: {e}")
    return table_cols

table_columns = get_table_columns()
print(f"  ✅ Table columns loaded: {list(table_columns.keys())}")


# ══════════════════════════════════════════════════════════════
# DATE INJECTION SYSTEM
# ══════════════════════════════════════════════════════════════

MONTH_NAMES = {
    'january': 1, 'february': 2, 'march': 3, 'april': 4,
    'may': 5, 'june': 6, 'july': 7, 'august': 8,
    'september': 9, 'october': 10, 'november': 11, 'december': 12,
    'jan': 1, 'feb': 2, 'mar': 3, 'apr': 4,
    'jun': 6, 'jul': 7, 'aug': 8, 'sep': 9,
    'sept': 9, 'oct': 10, 'nov': 11, 'dec': 12
}
QUARTER_MONTHS = {1: 1, 2: 4, 3: 7, 4: 10}

def get_quarter_range(year, quarter):
    start_month = QUARTER_MONTHS[quarter]
    start = datetime(year, start_month, 1)
    if quarter == 4:
        end = datetime(year + 1, 1, 1)
    else:
        end = datetime(year, QUARTER_MONTHS[quarter + 1], 1)
    return start.strftime('%Y-%m-%d'), end.strftime('%Y-%m-%d')

def detect_date_context(question):
    q = question.lower().strip()
    today = datetime.now()
    now_year = today.year
    now_month = today.month
    result = None

    m = re.search(r'\b(\d{1,2})(?:st|nd|rd|th)?\s+(january|february|march|april|may|june|july|august|september|october|november|december|jan|feb|mar|apr|jun|jul|aug|sep|sept|oct|nov|dec)\s+(\d{4})\b', q)
    if m:
        day, month_str, year = int(m.group(1)), m.group(2), int(m.group(3))
        month = MONTH_NAMES[month_str]
        date_str = f"{year}-{month:02d}-{day:02d}"
        result = {'filter': f"order_date = '{date_str}'::date", 'desc': f"on {date_str}"}

    if not result:
        m = re.search(r'\b(january|february|march|april|may|june|july|august|september|october|november|december|jan|feb|mar|apr|jun|jul|aug|sep|sept|oct|nov|dec)\s+(\d{1,2})(?:st|nd|rd|th)?\s+(\d{4})\b', q)
        if m:
            month_str, day, year = m.group(1), int(m.group(2)), int(m.group(3))
            month = MONTH_NAMES[month_str]
            date_str = f"{year}-{month:02d}-{day:02d}"
            result = {'filter': f"order_date = '{date_str}'::date", 'desc': f"on {date_str}"}

    if not result:
        m = re.search(r'\b(\d{4}-\d{2}-\d{2})\b', q)
        if m:
            date_str = m.group(1)
            result = {'filter': f"order_date = '{date_str}'::date", 'desc': f"on {date_str}"}

    if not result:
        m = re.search(r'\b(january|february|march|april|may|june|july|august|september|october|november|december|jan|feb|mar|apr|jun|jul|aug|sep|sept|oct|nov|dec)\s+(\d{4})\b', q)
        if m:
            month_str, year = m.group(1), int(m.group(2))
            month = MONTH_NAMES[month_str]
            start = f"{year}-{month:02d}-01"
            end_month = month + 1 if month < 12 else 1
            end_year = year if month < 12 else year + 1
            end = f"{end_year}-{end_month:02d}-01"
            result = {'filter': f"order_date >= '{start}'::date AND order_date < '{end}'::date", 'desc': f"in {m.group(1).title()} {year}"}

    if not result:
        m = re.search(r'\bq([1-4])\s+(?:of\s+)?(\d{4})\b', q)
        if m:
            quarter, year = int(m.group(1)), int(m.group(2))
            start, end = get_quarter_range(year, quarter)
            result = {'filter': f"order_date >= '{start}'::date AND order_date < '{end}'::date", 'desc': f"Q{quarter} {year}"}

    if not result:
        quarter_words = {'first': 1, 'second': 2, 'third': 3, 'fourth': 4}
        m = re.search(r'\b(first|second|third|fourth)\s+quarter\s+(?:of\s+)?(\d{4})\b', q)
        if m:
            quarter, year = quarter_words[m.group(1)], int(m.group(2))
            start, end = get_quarter_range(year, quarter)
            result = {'filter': f"order_date >= '{start}'::date AND order_date < '{end}'::date", 'desc': f"Q{quarter} {year}"}

    if not result:
        m = re.search(r'\b(in|during|for|of)\s+(202[0-9])\b', q)
        if not m:
            m = re.search(r'\b(202[0-9])\b', q)
            if m:
                year = int(m.group(1))
                result = {'filter': f"order_date >= '{year}-01-01'::date AND order_date < '{year+1}-01-01'::date", 'desc': f"in {year}"}
        else:
            year = int(m.group(2))
            result = {'filter': f"order_date >= '{year}-01-01'::date AND order_date < '{year+1}-01-01'::date", 'desc': f"in {year}"}

    if not result and re.search(r'\blast\s+week\b', q):
        start = (today - timedelta(days=7)).strftime('%Y-%m-%d')
        result = {'filter': f"order_date >= '{start}'::date AND order_date < CURRENT_DATE", 'desc': "last week"}

    if not result and re.search(r'\bthis\s+week\b', q):
        start_of_week = (today - timedelta(days=today.weekday())).strftime('%Y-%m-%d')
        result = {'filter': f"order_date >= '{start_of_week}'::date AND order_date < CURRENT_DATE", 'desc': "this week"}

    if not result and re.search(r'\blast\s+month\b', q):
        if now_month == 1:
            start = f"{now_year - 1}-12-01"
            end = f"{now_year}-01-01"
        else:
            start = f"{now_year}-{now_month-1:02d}-01"
            end = f"{now_year}-{now_month:02d}-01"
        result = {'filter': f"order_date >= '{start}'::date AND order_date < '{end}'::date", 'desc': "last month"}

    if not result and re.search(r'\bmonth\s+before\s+(last|that)\b|\b2\s+months\s+ago\b', q):
        two_months_ago = now_month - 2
        year_adj = now_year
        if two_months_ago <= 0:
            two_months_ago += 12
            year_adj -= 1
        one_month_ago = now_month - 1
        year_adj2 = now_year
        if one_month_ago <= 0:
            one_month_ago += 12
            year_adj2 -= 1
        start = f"{year_adj}-{two_months_ago:02d}-01"
        end = f"{year_adj2}-{one_month_ago:02d}-01"
        result = {'filter': f"order_date >= '{start}'::date AND order_date < '{end}'::date", 'desc': "month before last"}

    if not result and re.search(r'\bthis\s+month\b', q):
        start = f"{now_year}-{now_month:02d}-01"
        result = {'filter': f"order_date >= '{start}'::date AND order_date < CURRENT_DATE", 'desc': "this month"}

    if not result and re.search(r'\blast\s+quarter\b', q):
        current_quarter = (now_month - 1) // 3 + 1
        last_quarter = current_quarter - 1
        year = now_year
        if last_quarter == 0:
            last_quarter = 4
            year -= 1
        start, end = get_quarter_range(year, last_quarter)
        result = {'filter': f"order_date >= '{start}'::date AND order_date < '{end}'::date", 'desc': "last quarter"}

    if not result and re.search(r'\bthis\s+quarter\b', q):
        current_quarter = (now_month - 1) // 3 + 1
        start, end = get_quarter_range(now_year, current_quarter)
        result = {'filter': f"order_date >= '{start}'::date AND order_date < CURRENT_DATE", 'desc': "this quarter"}

    if not result and re.search(r'\blast\s+year\b', q):
        start = f"{now_year - 1}-01-01"
        end = f"{now_year}-01-01"
        result = {'filter': f"order_date >= '{start}'::date AND order_date < '{end}'::date", 'desc': "last year"}

    if not result and re.search(r'\bthis\s+year\b', q):
        start = f"{now_year}-01-01"
        result = {'filter': f"order_date >= '{start}'::date AND order_date < CURRENT_DATE", 'desc': "this year"}

    if not result:
        m = re.search(r'\blast\s+(\d+)\s+days?\b', q)
        if m:
            n = int(m.group(1))
            start = (today - timedelta(days=n)).strftime('%Y-%m-%d')
            result = {'filter': f"order_date >= '{start}'::date AND order_date < CURRENT_DATE", 'desc': f"last {n} days"}

    if not result:
        m = re.search(r'\blast\s+(\d+)\s+months?\b', q)
        if m:
            n = int(m.group(1))
            month = now_month - n
            year = now_year
            while month <= 0:
                month += 12
                year -= 1
            start = f"{year}-{month:02d}-01"
            result = {'filter': f"order_date >= '{start}'::date AND order_date < CURRENT_DATE", 'desc': f"last {n} months"}

    if not result:
        m = re.search(r'\bbetween\s+(january|february|march|april|may|june|july|august|september|october|november|december|jan|feb|mar|apr|jun|jul|aug|sep|sept|oct|nov|dec)\s+(\d{4})\s+and\s+(january|february|march|april|may|june|july|august|september|october|november|december|jan|feb|mar|apr|jun|jul|aug|sep|sept|oct|nov|dec)\s+(\d{4})\b', q)
        if m:
            m1, y1 = MONTH_NAMES[m.group(1)], int(m.group(2))
            m2, y2 = MONTH_NAMES[m.group(3)], int(m.group(4))
            start = f"{y1}-{m1:02d}-01"
            end_month = m2 + 1 if m2 < 12 else 1
            end_year = y2 if m2 < 12 else y2 + 1
            end = f"{end_year}-{end_month:02d}-01"
            result = {'filter': f"order_date >= '{start}'::date AND order_date < '{end}'::date", 'desc': f"between {m.group(1).title()} {y1} and {m.group(3).title()} {y2}"}

    return result


def inject_date_into_prompt(question):
    comparison_words = [
        "compare", "vs", "versus", "month before", "last month vs",
        "this year vs", "last year vs", "previous", "month before last",
        "2 months ago", "last quarter vs", "this quarter vs"
    ]
    if any(word in question.lower() for word in comparison_words):
        print(f"  📅 Comparison question — skipping date injection")
        return None
    date_ctx = detect_date_context(question)
    if date_ctx:
        print(f"  📅 Date detected: {date_ctx['desc']} → {date_ctx['filter']}")
        return date_ctx
    return None


# ══════════════════════════════════════════════════════════════
# INTENT ROUTER
# Classifies every question before processing
# Uses rule-based checks first (fast), then AI for ambiguous ones
# ══════════════════════════════════════════════════════════════

GREETING_EXACT = {
    "hello", "hi", "hey", "good morning", "good afternoon",
    "good evening", "how are you", "thanks", "thank you",
    "bye", "goodbye", "see you", "thats all", "that's all",
    "good job", "well done", "nice work", "what can you do"
}

CONTEXTUAL_REFS = {
    "that", "it", "its", "those", "they", "them", "this",
    "that region", "that category", "that customer", "that product",
    "that state", "that city", "same", "also", "too", "as well",
    "more about", "tell me more", "explain", "elaborate", "details","first", "second", "third", "fourth", "fifth",
    "first one", "second one", "top one", "number one",
    "that person", "that customer", "he", "she", "they"
}

RECOMMENDATION_KEYWORDS = {
    "should", "recommend", "suggest", "ideal", "best strategy",
    "what would you", "advice", "advise", "optimize", "improve",
    "increase", "decrease", "focus on", "invest", "priority",
    "inventory", "reorder", "stock", "forecast", "predict",
    "what can we do", "how can we", 
    "how to improve", "action", "plan"
}

SUBJECTIVE_SQL_KEYWORDS = {
    "why", "explain why", "what caused", "reason for",
    "what drives", "insight", "analyse", "analyze",
    "why is", "why are", "why does", "performing better",
    "performing worse", "outperforming", "underperforming"
}

def classify_intent(question):
    """
    Classifies into: generic | sql | recommendation | subjective_sql | contextual
    Hard rules for clear cases, AI for everything else.
    """
    q = question.lower().strip().rstrip("!?.")

    # Rule 1: Exact greetings
    if q in GREETING_EXACT:
        return "generic"
    
    # Get real context (skip summary pseudo-entry)
    # Get real context (skip summary pseudo-entry)
    real_ctx = [ex for ex in context if ex.get("question") != "__summary__"]

    # If context exists — always contextual, no exceptions
    if real_ctx:
        return "contextual"

    # No context — first question, normal routing
    sql_phrases = [
        "how many", "how much", "total sales", "total profit", "total orders",
        "total returns", "show me", "give me all details", "all details",
        "details of", "which city", "which state", "which region",
        "which category", "which segment", "which channel", "which device",
        "which month", "which year", "which quarter", "which payment",
        "top ", "bottom ", "highest", "lowest", "most", "least",
        "average delivery", "average order"
    ]
    if any(p in q for p in sql_phrases):
        return "sql"

    return "sql"

# ════════════════════════════════════════════ ══════════════════
# HANDLERS — one per intent type
# ══════════════════════════════════════════════════════════════

def handle_generic(question):
    """Handles greetings and casual conversation using Mistral."""
    prompt = f"""You are a friendly Sales Analytics Assistant for Superstore.
The user said: '{question}'
Respond naturally in 1-2 sentences. Be warm and helpful.
If they greet you, greet back and offer to help with sales data.
Never mention SQL or technical terms. Never start with 'I'."""
    return call_model(prompt, model=FAST_MODEL, temperature=0.8, num_ctx=512, timeout=30)


def _resolve_reference(question, prev_result_str):
    """
    Parse the previous raw result and resolve vague references like
    'lowest', 'highest', 'first', 'second', 'that one' into exact values.
    Returns a dict with resolved info, or empty dict if can't resolve.
    """
    if not prev_result_str:
        return {}
    try:
        import pandas as pd
        import io
        df = pd.read_csv(io.StringIO(prev_result_str), sep=r'\s{2,}', engine='python')
        df = df.dropna(how='all')
        if df.empty:
            return {}

        q = question.lower()
        resolved = {}

        # Detect ordinal references
        ordinal_map = {
            'first': 0, '1st': 0, 'number one': 0, 'top one': 0,
            'second': 1, '2nd': 1, 'number two': 1,
            'third': 2, '3rd': 2, 'number three': 2,
            'fourth': 3, '4th': 3,
            'fifth': 4, '5th': 4,
        }
        for word, idx in ordinal_map.items():
            if word in q and idx < len(df):
                row = df.iloc[idx]
                resolved['row'] = row.to_dict()
                resolved['description'] = f"the {word} row: {row.to_dict()}"
                return resolved

        # Detect lowest/minimum reference
        if any(w in q for w in ['lowest', 'minimum', 'least', 'bottom', 'smallest', 'worst']):
            # Find numeric columns
            num_cols = df.select_dtypes(include='number').columns.tolist()
            if num_cols:
                col = num_cols[-1]  # last numeric col (usually the metric)
                idx = df[col].idxmin()
                row = df.loc[idx]
                resolved['row'] = row.to_dict()
                resolved['description'] = f"the row with the LOWEST {col} value ({row[col]}): {row.to_dict()}"
                return resolved

        # Detect highest/maximum reference
        if any(w in q for w in ['highest', 'maximum', 'most', 'top', 'best', 'largest']):
            num_cols = df.select_dtypes(include='number').columns.tolist()
            if num_cols:
                col = num_cols[-1]
                idx = df[col].idxmax()
                row = df.loc[idx]
                resolved['row'] = row.to_dict()
                resolved['description'] = f"the row with the HIGHEST {col} value ({row[col]}): {row.to_dict()}"
                return resolved

        return {}
    except Exception as e:
        print(f"  ⚠️ _resolve_reference error: {e}")
        return {}


def _classify_contextual_subintent(question, prev_question, prev_answer, prev_result):
    q = question.lower().strip()

    # Subjective — why/explain questions
    if any(w in q for w in [
        "why", "explain", "what caused", "reason", "what drives",
        "how come", "what happened", "insight", "analyse", "analyze"
    ]):
        return "subjective"

    # Text only — no DB needed
    if any(w in q for w in [
        "bundle", "group them", "categorise", "categorize",
        "summarise", "summarize", "list them", "can u list",
        "show them", "what are they", "name them"
    ]):
        return "text"

    # Everything else — go to DB
    return "sql_lookup"
    """
    Decides how to handle a contextual follow-up:
    - sql_lookup  : needs a fresh DB query referencing previous result
    - subjective  : asks why/explain/interpret previous answer
    - text        : pure text follow-up (bundle, list, summarise, etc.)
    """
    prompt = f"""You are classifying a follow-up question for a sales chatbot.

Previous question: '{prev_question}'
Previous answer summary: '{prev_answer[:400]}'
Previous raw data: '{prev_result[:400]}'

Follow-up question: '{question}'

Classify into ONE of:
- sql_lookup  : question needs a new database query using specific values from previous answer
                (e.g. "give details on the lowest", "show orders for that customer", 
                "what did that region buy", "more about the first one", "compare those two")
- subjective  : question asks why something happened or needs business interpretation
                (e.g. "why did that happen", "what caused this", "explain the drop")
- text        : question can be answered from the previous answer text alone, no DB needed
                (e.g. "bundle those products", "list them by category", "summarise that")

Reply with ONE word only: sql_lookup, subjective, or text"""

    result = call_model(prompt, model=FAST_MODEL, temperature=0.1, num_ctx=512, timeout=60)
    print(f"  🔍 Sub-intent raw result: '{result}'")
    result = result.strip().lower().split()[0] if result.strip() else "text"
    return result if result in {"sql_lookup", "subjective", "text"} else "text"


def handle_contextual(question, ctx):
    """
    Handles follow-up questions with 3 sub-intents:
    - sql_lookup  : re-queries DB using previous result as reference
    - subjective  : interprets previous answer with Mistral
    - text        : pure text response using previous answer
    """
    # Filter out summary pseudo-entries
    real_ctx = [ex for ex in ctx if ex.get("question") != "__summary__"]
    summary_entry = next((ex for ex in ctx if ex.get("question") == "__summary__"), None)

    if not real_ctx:
        return handle_generic(question)

    last = real_ctx[-1]
    prev_question = last['question']
    prev_answer = last['answer']
    prev_result = last.get('result', '')

    sub = _classify_contextual_subintent(question, prev_question, prev_answer, prev_result)
    print(f"  🔀 Contextual sub-intent: {sub}")

    # ── SQL LOOKUP ─────────────────────────────────────────────
    if sub == "sql_lookup":
        # Step 1: Resolve vague references (lowest, first, that one) in Python
        resolved = _resolve_reference(question, prev_result)
        resolved_desc = resolved.get('description', '')
        resolved_row = resolved.get('row', {})

        # Step 2: Extract the most likely entity name/key from resolved row
        # so Gemma3 gets an exact WHERE clause anchor
        entity_hint = ""
        for key in ['customer_name', 'region', 'state', 'city', 'category',
                    'sub_category', 'product_name', 'segment', 'ship_mode']:
            if key in resolved_row:
                entity_hint = f"The specific {key} to query is: '{resolved_row[key]}'"
                break

        enriched = f"""{question}

CONTEXT:
- Previous question: '{prev_question}'
- Previous answer: {prev_answer[:400]}
- Previous raw data:
{prev_result[:600]}

RESOLVED REFERENCE:
{resolved_desc if resolved_desc else "Could not auto-resolve — use the raw data above to infer."}
{entity_hint}

CRITICAL INSTRUCTIONS:
- Do NOT write a query that finds the global minimum/maximum from the whole table
- The user is referring to a specific row from the PREVIOUS result above
- Use the resolved entity name/value above as the WHERE filter
- For example if previous result had top 5 customers and user says 'lowest', 
  query for the customer with the lowest value IN THAT LIST, not the whole DB
- Write a detailed follow-up query about that specific entity"""

        date_ctx = inject_date_into_prompt(question)
        sql = text_to_sql(enriched, date_ctx)
        print(f"  ├─ Contextual SQL: {sql}")

        if not sql or "UNRELATED" in sql.upper() or not sql.strip().upper().startswith(("SELECT", "WITH")):
            sub = "text"
        else:
            result = run_query(sql, question)
            if result is None or result.empty:
                sub = "text"
            else:
                if needs_visualisation(question):
                    create_visualisation(result, question)
                # Summarise with context of what was resolved
                context_hint = f"Note: The user asked about {resolved_desc}" if resolved_desc else ""
                answer = summarise(f"{question} ({context_hint})", result)
                context.append({
                    "question": question,
                    "sql": sql,
                    "answer": answer,
                    "result": result.to_string()
                })
                if len(context) > 10:
                    context.pop(0)
                return answer

    # ── SUBJECTIVE ─────────────────────────────────────────────
    if sub == "subjective":
        summary_text = summary_entry['answer'] if summary_entry else ""
        prompt = f"""You are an expert Sales Analytics Consultant for Superstore.

{f'Conversation summary so far: {summary_text}' if summary_text else ''}

Previous question: '{prev_question}'
Previous answer: {prev_answer[:600]}
Raw data: {prev_result[:400]}

User now asks: '{question}'

Provide a short analytical interpretation (2-3 sentences max).
Answer the 'why' or 'what this means' using the data above.
Be specific — reference actual numbers/names from the data. Never start with 'I'."""
        answer = call_model(prompt, model=FAST_MODEL, temperature=0.5, num_ctx=1024, timeout=45)
        context.append({
            "question": question,
            "sql": "",
            "answer": answer,
            "result": ""
        })
        if len(context) > 10:
            context.pop(0)
        return answer

    # ── TEXT (default) ─────────────────────────────────────────
    summary_text = summary_entry['answer'] if summary_entry else ""
    prompt = f"""You are a Sales Analytics Assistant for Superstore.

{f'Conversation summary so far: {summary_text}' if summary_text else ''}

Previous question: '{prev_question}'
Previous answer:
{prev_answer}

User now asks: '{question}'

STRICT RULES:
- ONLY use data explicitly listed in the previous answer above
- Do NOT invent products, names, or numbers not in the previous answer
- If asked to bundle, group ONLY items from the previous answer by their category
- If the previous answer has no relevant data, say so clearly
- Never start with 'I'
- Be concise"""

    answer = call_model(prompt, model=FAST_MODEL, temperature=0.4, num_ctx=1024, timeout=45)
    if not answer.strip():
        answer = "I couldn't generate a response. Please try rephrasing."
    context.append({
        "question": question,
        "sql": "",
        "answer": answer,
        "result": ""
    })
    if len(context) > 10:
        context.pop(0)
    return answer


def handle_recommendation(question, ctx):
    """
    Handles recommendation questions like Amazon.
    Top 2 products from each of 5 categories = 10 products total.
    """
    # If question references previous answer, redirect to contextual
    context_refs = [
        "above", "those", "from the list", "suggested", "the ones", "bundle",
        "from that", "of those", "these products", "the above", "from above",
        "mentioned", "listed", "previous", "from the above"
    ]
    if ctx and any(ref in question.lower() for ref in context_refs):
        return handle_contextual(question, ctx), ctx[-1].get('result', '')

    # Build WHERE clause from question
    q = question.lower()
    filters = []

    age_match = re.search(r'(\d+)\s*[-–to]+\s*(\d+)', q)
    if age_match:
        filters.append(f"c.age BETWEEN {age_match.group(1)} AND {age_match.group(2)}")
    elif 'young' in q:
        filters.append("c.age BETWEEN 21 AND 30")
    elif 'middle aged' in q or 'middle-aged' in q:
        filters.append("c.age BETWEEN 41 AND 50")
    elif 'senior' in q or 'elderly' in q:
        filters.append("c.age > 50")
    elif 'teenager' in q or 'teen' in q:
        filters.append("c.age BETWEEN 11 AND 20")

    if 'non-prime' in q or 'non prime' in q:
        filters.append("c.prime_status = 'Non-Prime'")
    elif 'prime' in q:
        filters.append("c.prime_status = 'Prime'")

    if 'mobile' in q and 'phone' not in q:
        filters.append("c.channel = 'Mobile'")
    elif 'web' in q:
        filters.append("c.channel = 'Web'")

    if 'android' in q:
        filters.append("c.device = 'Android'")
    elif 'ios' in q:
        filters.append("c.device = 'iOS'")

    where_clause = "WHERE " + " AND ".join(filters) if filters else ""

    # Hardcoded SQL — top 2 per category guaranteed
    sql = f"""
WITH ranked AS (
    SELECT
        o.product_name,
        o.category,
        COUNT(DISTINCT o.order_id) AS purchase_count,
        ROW_NUMBER() OVER (
            PARTITION BY o.category
            ORDER BY COUNT(DISTINCT o.order_id) DESC
        ) AS rn
    FROM orders o
    JOIN customers c ON o.customer_name = c.customer_name
    {where_clause}
    GROUP BY o.product_name, o.category
)
SELECT product_name, category, purchase_count
FROM ranked
WHERE rn <= 2
ORDER BY category, rn
"""

    relevant_data_str = ""
    product_list_from_db = []

    try:
        result = run_query(sql, question)
        if result is not None and not result.empty and 'product_name' in result.columns:
            relevant_data_str = result.to_string()
            product_list_from_db = result['product_name'].tolist()
    except Exception as e:
        print(f"  ⚠️ Recommendation SQL error: {e}")

    if product_list_from_db:
        if age_match:
            group_desc = f"customers aged {age_match.group(1)}-{age_match.group(2)}"
        elif 'non-prime' in q or 'non prime' in q:
            group_desc = "Non-Prime members"
        elif 'prime' in q:
            group_desc = "Prime members"
        elif 'android' in q:
            group_desc = "Android users"
        elif 'ios' in q:
            group_desc = "iOS users"
        elif 'mobile' in q and 'phone' not in q:
            group_desc = "Mobile channel customers"
        elif 'web' in q:
            group_desc = "Web channel customers"
        else:
            group_desc = "customers matching your criteria"

        header = f"Here are the recommended products, which others from {group_desc} have also bought:\n"
        numbered = "\n".join([f"{i+1}. {p}" for i, p in enumerate(product_list_from_db)])
        why_prompt = f"In one sentence under 20 words, explain why these products suit {group_desc}. Start with: These are the top products purchased by"
        why = call_model(why_prompt, model=FAST_MODEL, temperature=0.3, num_ctx=256, timeout=30)
        why = why.strip().split('\n')[0]
        answer = f"{header}\n{numbered}\n\n{why}"
        return answer, relevant_data_str

    # Fallback with your original rules
    prompt = f"""You are a product recommendation engine for a retail company.

User asks: '{question}'

RECOMMENDATION RULES:
- recommend 10 products listed one by one
- take products which other people of the same category have bought
- involve products from at least 3 categories
- make sure 10 products are divided across categories, not just 1 or 2
- output format:
Here are the recommended products, which others also from the same category have bought (mention the category like age 20-30)
1. product name
2. product name
... till 10
At the end: one liner explanation of why these products.

ALLOWED PRODUCTS ONLY — never use anything outside this list:
Electronics: Samsung Galaxy Phone, iPhone, OnePlus Phone, Redmi Phone, Realme Phone, Dell Laptop, HP Laptop, Lenovo Laptop, Asus Laptop, MacBook, Bluetooth Speaker, Earphones, Headphones, Smart Watch
Furniture: Office Chair, Study Chair, Dining Chair, Recliner Chair, Office Desk, Study Table, Coffee Table, Side Table, Bookshelf, Cabinet, Wardrobe, Drawer Unit, TV Stand, Shoe Rack
Office Supplies: Pen, Pencil, Marker, Highlighter, Notebook, Register, Printer Paper, Sticky Notes, Stapler, Tape Dispenser, File Folder, Envelope, Calculator, Clipboard
Home & Kitchen: Mixer Grinder, Electric Kettle, Toaster, Rice Cooker, Water Bottle, Lunch Box, Storage Container, Dinner Set, Wall Clock, Table Lamp, Curtain, Bedsheet, Pillow, Mirror
Fashion: T-Shirt, Shirt, Jeans, Trousers, Dress, Jacket, Sneakers, Formal Shoes, Sandals, Sports Shoes, Backpack, Handbag, Wallet, Belt"""

    answer = call_model(prompt, model=FAST_MODEL, temperature=0.4, num_ctx=4096, timeout=120)
    return answer, relevant_data_str

def handle_customer_recommendation(question, customer_name):
    
    

    prompt = f"""You are a sales assistant for Superstore.

            Based on the sales data below, list the 10 products popular with customers aged {age_group}.

            Data:
            {sql_result}

            Rules:
            - Output ONLY a numbered list of product names and 1 liner overall explanation of why these products
            - No descriptions, no explanations, no strategy, no expected outcomes
            - Just the product name on each line
            - Example format:
            why these products-1 liner explanation
            1. Staple Envelopes
            2. Easy-staple Paper
            3. KI Adjustable-Height Table

            Reply with the numbered list only."""
            
    try:
        sql = f"""
WITH customer_profile AS (
    SELECT c.age, c.sex, c.prime_status
    FROM customers c
    WHERE LOWER(c.customer_name) = LOWER('{customer_name}')
),
customer_products AS (
    SELECT DISTINCT o.product_name, o.category
    FROM orders o
    WHERE LOWER(o.customer_name) = LOWER('{customer_name}')
),
similar_customers AS (
    SELECT c.customer_name
    FROM customers c, customer_profile sp
    WHERE c.sex = sp.sex
    AND c.age BETWEEN sp.age - 5 AND sp.age + 5
    AND c.prime_status = sp.prime_status
    AND LOWER(c.customer_name) != LOWER('{customer_name}')
),
recommended_products AS (
    SELECT 
        o.product_name,
        o.category,
        COUNT(DISTINCT o.customer_name) AS bought_by,
        ROUND(AVG(o.sales)::numeric, 2) AS avg_price
    FROM orders o
    WHERE o.customer_name IN (SELECT customer_name FROM similar_customers)
    AND o.product_name NOT IN (SELECT product_name FROM customer_products)
    GROUP BY o.product_name, o.category
    ORDER BY bought_by DESC
    LIMIT 10
)
SELECT 
    r.product_name,
    r.category,
    r.bought_by AS similar_customers_bought,
    r.avg_price,
    cp.age,
    cp.sex,
    cp.prime_status
FROM recommended_products r, customer_profile cp
"""
        result = run_query(sql, question)

        if result is None or result.empty:
            return f"Couldn't find recommendations for {customer_name}. They may not be in our customer database."

        # Get customer profile for context
        age = result['age'].iloc[0]
        sex = result['sex'].iloc[0]
        prime = result['prime_status'].iloc[0]

        # Build recommendation text using Mistral
        products_str = result[['product_name', 'category', 'similar_customers_bought', 'avg_price']].to_string(index=False)

        prompt = f"""You are a retail recommendation engine for Superstore.

Customer: {customer_name}
Profile: {sex}, Age {age}, {prime} member

These products are bought by similar customers ({sex}, similar age, same prime status) 
that {customer_name} has NOT purchased yet, ranked by popularity:

{products_str}

RECOMMENDATION RULES:
- when asked to recommend products based on age group or just a individual name of a customer or based on sex or region or etc, recommend 10 products.
- the products have to be listed one by one.
- take the products which other people of the same category have used as for eg if the question is recommend products for age group 20-30, recommend products which other people have used from age 20-30 and involve products from atleast 3 categories.
- make sure the 10 products which you will list are divided across 3 categories and are not of one or 2 product categories.
- output should be like:
Here are the recommended products, which others also from the same category have bought(mention what category like age 20-30)
1.product name
2. product name
till 10

at the end one liner explanation of this as in why these products"""


        answer = call_model(prompt, model=FAST_MODEL, temperature=0.6,
                           num_ctx=2048, timeout=60)

        # Save to context
        context.append({
            "question": question,
            "sql": sql,
            "answer": answer,
            "result": result.to_string()
        })
        if len(context) > 10:
            context.pop(0)

        return answer

    except Exception as e:
        print(f"  ❌ Recommendation error: {e}")
        return f"Could not generate recommendations for {customer_name}. Please try again."

def handle_subjective_sql(question, ctx):
    """
    Handles questions needing data AND interpretation.
    Gemma3 gets the data, Mistral interprets it.
    """
    # Step 1: Get data with Gemma3
    date_ctx = inject_date_into_prompt(question)
    sql = text_to_sql(question, date_ctx)
    print(f"  ├─ SQL: {sql}")

    data_str = ""
    result = None

    if sql and not "UNRELATED" in sql.upper() and sql.strip().upper().startswith(("SELECT", "WITH")):
        result = run_query(sql, question)
        if result is not None and not result.empty:
            data_str = result.head(20).to_string()
            if needs_visualisation(question):
                create_visualisation(result, question)

    # Step 2: Interpret with Mistral
    context_str = ""
    if ctx:
        context_str = f"\nPrevious context: {ctx[-1]['question']} → {ctx[-1]['answer']}"

    prompt = f"""You are an expert Sales Analytics Consultant for Superstore.

The user asks: '{question}'
Database results: {data_str if data_str else 'No specific data found'}
{context_str}

Provide an insightful interpretation and explanation. Answer the 'why' behind the numbers.
Be analytical and business-focused. Max 2 sentences. Never start with 'I'."""

    answer = call_model(prompt, model=FAST_MODEL, temperature=0.6,
                       num_ctx=4096, timeout=120)

    # Save to context
    if result is not None and not result.empty:
        context.append({
            "question": question,
            "sql": sql if sql else "",
            "answer": answer,
            "result": data_str
        })
        if len(context) > 10:
            context.pop(0)

    return answer


def handle_sql(question):
    """
    Handles data questions: generates SQL with Gemma3, summarises with Mistral.
    """
    global latest_chart_buf

    date_ctx = inject_date_into_prompt(question)
    sql = text_to_sql(question, date_ctx)
    print(f"  ├─ SQL: {sql}")

    if "UNRELATED" in sql.upper():
        # Check if question contains a name — retry with explicit customer hint
        name_hint_phrases = ["details of", "tell me about", "info on", "all details"]
        if any(phrase in question.lower() for phrase in name_hint_phrases):
            retry_q = f"{question} — look up this person in the customers table and orders table"
            sql = text_to_sql(retry_q, date_ctx)
            if "UNRELATED" in sql.upper() or not sql.strip().upper().startswith(("SELECT", "WITH")):
                return "Couldn't find details for that person. Please check the name and try again."
        elif context and len(question.split()) <= 6:
            retry_q = f"{question} (referring to previous answer about {context[-1]['question']})"
            sql = text_to_sql(retry_q, date_ctx)
            if "UNRELATED" in sql.upper() or not sql.strip().upper().startswith(("SELECT", "WITH")):
                return handle_generic(question)
        else:
            return handle_generic(question)

    if not sql.strip().upper().startswith(("SELECT", "WITH")):
        return handle_generic(question)

    result = run_query(sql, question)

    # For product listing questions — return clean list, skip summarise
    listing_phrases = [
        "what products", "which products", "list products",
        "show products", "what items", "which items"
    ]
    if any(phrase in question.lower() for phrase in listing_phrases):
        if 'product_name' in result.columns:
            products = result['product_name'].tolist()
            answer = "\n".join([f"{i+1}. {p}" for i, p in enumerate(products)])
            context.append({"question": question, "sql": sql, "answer": answer, "result": result.to_string()})
            if len(context) > 10: context.pop(0)
            return answer

    if result is None or result.empty:
        return "Couldn't find data for that. Try rephrasing — for example: 'total sales last month' or 'profit by region'."

    detail_phrases = ["all orders", "all data", "give me all", "give me details", "show me details"]
    is_person_query = any(phrase in question.lower() for phrase in 
                        ["of ", "for ", "about ", "details of", "data of", "data for"])

    if any(phrase in question.lower() for phrase in detail_phrases) and not is_person_query:
        answer = result.to_string(index=False)
        context.append({"question": question, "sql": sql, "answer": answer, "result": result.to_string()})
        if len(context) > 10: context.pop(0)
        return answer

    if needs_visualisation(question):
        create_visualisation(result, question)

    answer = summarise(question, result)

    context.append({
        "question": question,
        "sql": sql,
        "answer": answer,
        "result": result.to_string() if result is not None else ""
    })
    if len(context) > 10:
        context.pop(0)

    return answer


# ══════════════════════════════════════════════════════════════
# SQL GENERATION — text_to_sql()
# Uses Gemma3 (SQL_MODEL) — best model for SQL
# ══════════════════════════════════════════════════════════════

def text_to_sql(question, date_ctx=None):
    context_str = ""
    if context:
        context_str = "\n\nCONVERSATION HISTORY:\n"
        for exchange in context:
            context_str += f"User: {exchange['question']}\n"
            context_str += f"SQL: {exchange['sql']}\n"
            context_str += f"Answer: {exchange['answer']}\n"
            context_str += f"Result data: {exchange.get('result', '')}\n\n"
        context_str += """CRITICAL FOLLOW UP RULES:
- Read the conversation history AND result data above very carefully
- The result data shows the EXACT values returned by the database
- When user says 'that region' or 'it' or 'its' — look at the previous answer to find the exact region name and add WHERE region = 'exact name'
- When user says 'that customer' — add WHERE customer_name = 'exact name from previous answer'
- When user says 'that category' — add WHERE category = 'exact name from previous answer'
- When user says 'that city' — add WHERE city = 'exact name from previous answer'
- When user says 'that state' — add WHERE state = 'exact name from previous answer'
- NEVER use WHERE customer_name when the previous question was about a region
- NEVER use WHERE region when the previous question was about a customer
- Always match the filter column to what was being discussed in the previous question
- For 'which has the lowest' after a highest question — use same GROUP BY and metric but ORDER BY ASC LIMIT 1
- Never return UNRELATED when there is conversation history\n"""

    date_instruction = ""
    if date_ctx:
        date_instruction = f"""
CRITICAL DATE INSTRUCTION:
The user is asking about: {date_ctx['desc']}
Use EXACTLY this date filter in your WHERE clause (already calculated for you):
  {date_ctx['filter']}
Do NOT calculate dates yourself — just use the filter above as-is.
"""

    prompt = f"""You are an expert Sales Analytics Assistant for a retail company. Convert plain English questions to valid PostgreSQL SQL queries.

Today's date is {datetime.now().strftime('%Y-%m-%d')}.

{metadata}

The orders table has data from 2022 to 2026.
The returns table is linked to orders via order_id.
The customers table has demographics: age, sex, prime_status, payment_mode, channel, device.
The customers table links to orders via customer_name.

MOST IMPORTANT RULE:
- If the question is unrelated to the database return exactly: UNRELATED
- NEVER return SQL comments (lines starting with --)
- NEVER return explanations instead of SQL

{date_instruction}

CUSTOMERS TABLE RULES:
- customers table has: customer_name, age, sex, prime_status, payment_mode, channel, device
- To join: JOIN customers c ON o.customer_name = c.customer_name
- For age group questions: WHERE c.age BETWEEN 20 AND 30
- For prime status: WHERE c.prime_status = 'Prime' or 'Non-Prime'
- For channel: WHERE c.channel = 'Web' or 'Mobile'
- For device: WHERE c.device = 'Android' or 'iOS' or 'Desktop'
- Always prefix customers columns with c. alias
- For "details of [name]" or "all details of [name]": 
  SELECT c.*, o.* FROM customers c 
  LEFT JOIN orders o ON c.customer_name = o.customer_name 
  WHERE LOWER(c.customer_name) = LOWER('name here') LIMIT 20
- NEVER return UNRELATED for questions containing a person's name
- If a name is mentioned, always look them up in customers or orders table
- customers table has NO customer_id column — always use COUNT(DISTINCT customer_name) to count customers
- Never use customer_id in any query
-- For age group questions, ALWAYS use this exact SQL structure:
  SELECT
    CASE 
      WHEN c.age BETWEEN 0 AND 10 THEN '0-10 (Child)'
      WHEN c.age BETWEEN 11 AND 20 THEN '11-20 (Teenager)'
      WHEN c.age BETWEEN 21 AND 30 THEN '21-30 (Young Adult)'
      WHEN c.age BETWEEN 31 AND 40 THEN '31-40 (Adult)'
      WHEN c.age BETWEEN 41 AND 50 THEN '41-50 (Middle Aged)'
      WHEN c.age BETWEEN 51 AND 60 THEN '51-60 (Senior)'
      WHEN c.age > 60 THEN '60+ (Elderly)'
      ELSE 'Unknown'
    END AS age_group,
    SUM(o.sales) AS total_sales
  FROM orders o
  JOIN customers c ON o.customer_name = c.customer_name
  GROUP BY age_group
  ORDER BY total_sales DESC
- The CASE statement goes inside SELECT, never after JOIN or WHERE
- GROUP BY age_group refers to the alias, not the full CASE expression
- When question asks for customer details WITH age and prime status, always include c.age and c.prime_status in SELECT
- Always add c.age, c.prime_status to GROUP BY when they are in SELECT
- Example: SELECT o.customer_name, ROUND(SUM(o.profit)::numeric, 2) AS total_profit, c.age, c.prime_status
  FROM orders o JOIN customers c ON o.customer_name = c.customer_name
  GROUP BY o.customer_name, c.age, c.prime_status
  ORDER BY total_profit DESC LIMIT 5

RELATIONSHIP BETWEEN TABLES:
- orders and returns are linked via order_id
- Always read the DATABASE STRUCTURE above carefully before writing SQL
- Only use columns that actually exist in each table as listed in DATABASE STRUCTURE
- For return rate by any group always use this exact pattern:
  SELECT o.group_column, ROUND((COUNT(DISTINCT r.return_id)::numeric / NULLIF(COUNT(DISTINCT o.order_id)::numeric, 0)) * 100, 2) AS return_rate FROM orders o LEFT JOIN returns r ON r.order_id = o.order_id GROUP BY o.group_column ORDER BY return_rate DESC
- Always prefix ambiguous columns with table alias
- Never write correlated subqueries for return rates — always use LEFT JOIN with COUNT
- The returns table does NOT have region, state, city, segment, category, customer_name or product columns
- To get returns by region: JOIN orders ON returns.order_id = orders.order_id and GROUP BY orders.region
- Any time user asks for returns breakdown BY any group — always JOIN returns with orders

BASIC SQL RULES:
- Return ONLY the raw SQL query, no backticks, no markdown, no explanations, no semicolon
- Always include FROM clause in every query
- Always include every non-aggregate column in GROUP BY
- Always use ROUND(SUM(column)::numeric, 2) for decimal results
- Never nest aggregate functions like SUM(SUM()) or ROUND(ROUND())
- Always alias columns with meaningful names using AS
- Always include the GROUP BY column in SELECT
- highest maximum top most all mean ORDER BY DESC
- lowest minimum bottom least all mean ORDER BY ASC
- When user asks for sales always use SUM(sales) never SUM(profit)
- When user asks for profit always use SUM(profit) never SUM(sales)
- Always use lowercase table names
- Never use COUNT(*) always use COUNT(DISTINCT column_name)
- When using JOIN always prefix ALL column references with table alias
- For division inside ROUND, ALWAYS cast the numerator: ROUND(SUM(o.sales)::numeric / NULLIF(COUNT(DISTINCT o.order_id), 0), 2)
- Never cast only the denominator — always cast the numerator or the whole expression
- Correct: ROUND(SUM(sales)::numeric / NULLIF(COUNT(DISTINCT order_id), 0), 2)
- Wrong:   ROUND(SUM(sales) / COUNT(DISTINCT order_id)::numeric, 2)

AGGREGATION RULES:
- For top N: ORDER BY value DESC LIMIT N
- For bottom N: ORDER BY value ASC LIMIT N
- For profit margin: ROUND((SUM(profit)::numeric / NULLIF(SUM(sales)::numeric, 0)) * 100, 2)
- For average order value: ROUND((SUM(sales)::numeric / NULLIF(COUNT(DISTINCT order_id), 0)), 2)
- For overall return rate use: SELECT ROUND(((SELECT COUNT(DISTINCT return_id) FROM returns)::numeric / NULLIF((SELECT COUNT(DISTINCT order_id) FROM orders)::numeric, 0)) * 100, 2) AS return_percentage
- For loss making orders: WHERE profit < 0
- For high discount orders: WHERE discount > 0.30
- For cumulative/running total by month use window function:
  SELECT TO_CHAR(order_date, 'YYYY-MM') AS year_month,
  ROUND(SUM(SUM(sales)) OVER (ORDER BY TO_CHAR(order_date, 'YYYY-MM'))::numeric, 2) AS cumulative_sales
  FROM orders GROUP BY TO_CHAR(order_date, 'YYYY-MM') ORDER BY year_month
- NEVER use regular SUM for cumulative — always use SUM(SUM()) OVER() window function
- For "highest discount" or "most discount" by city/region/state — always use AVG(discount) not SUM(discount)
- SUM(discount) is meaningless for discount analysis — always use AVG or MAX

DELIVERY RULES:
- For average delivery time: SELECT ROUND(AVG(EXTRACT(EPOCH FROM (ship_date - order_date)) / 86400)::numeric, 2) AS avg_delivery_days FROM orders
- For late deliveries: WHERE EXTRACT(EPOCH FROM (ship_date - order_date)) / 86400 > 5

DATE RULES:
- Today is {datetime.now().strftime('%Y-%m-%d')}. You ALWAYS know the current date.
- Always use DATE_TRUNC for date range comparisons
- For last month: WHERE order_date >= DATE_TRUNC('month', CURRENT_DATE - INTERVAL '1 months') AND order_date < DATE_TRUNC('month', CURRENT_DATE)
- For comparing two months always use CASE WHEN with DATE_TRUNC ranges
- NEVER use BETWEEN for date comparisons always use >= and <
- Always use INTERVAL '1 months' not INTERVAL '1 month'

LIMIT RULES:
- NEVER use LIMIT for breakdown questions: by category, by region, by segment, by state, by city, by year, by month, by quarter
- Only use LIMIT when user says: top N, bottom N, highest, lowest, most, least

RETURNS TABLE RULES:
- For refund analysis always use SUM(refund_amount) from returns table
- For return reasons use GROUP BY reason from returns table
- Always use LEFT JOIN returns with orders when calculating return rates by group

CTE RULES:
- When using CTEs with WITH clause, always reference the SUMMARY CTE not the base CTE
- If you have SeanReturns and SeanReturnSummary, use SeanReturnSummary in the final SELECT
- Never reference a base CTE in the final SELECT if a summary CTE exists for it
- Always include all CTEs in the FROM clause of the final SELECT

RECOMMENDATION RULES:
- when asked to recommend products based on age group or just a individual name of a customer or based on sex or region or etc, recommend 10 products.
- the products have to be listed one by one.
- take the products which other people of the same category have used as for eg if the question is recommend products for age group 20-30, recommend products which other people have used from age 20-30 and involve products from atleast 3 categories.
- make sure the 10 products which you will list are divided across 3 categories and are not of one or 2 product categories.
- output should be like:
Here are the recommended products, which others also from the same category have bought(mention what category like age 20-30)
1.product name
2. product name
till 10

at the end one liner explanation of this as in why these products.

RETURN TABKE RULES:
- For return rate by prime_status or any customer demographic, always use this exact pattern:
  SELECT 
    c.prime_status,
    ROUND((COUNT(DISTINCT r.return_id)::numeric / NULLIF(COUNT(DISTINCT o.order_id)::numeric, 0)) * 100, 2) AS return_rate
  FROM orders o
  JOIN customers c ON o.customer_name = c.customer_name
  LEFT JOIN returns r ON o.order_id = r.order_id
  GROUP BY c.prime_status
  ORDER BY return_rate DESC
- Always use COUNT(DISTINCT) for both return_id and order_id — never COUNT(*)
- Always LEFT JOIN returns so non-returned orders are included in the denominator

{context_str}

Current question: {question}"""

    # Use SQL_MODEL (gemma3) for SQL generation
    sql = call_model(prompt, model=SQL_MODEL, temperature=0.1,
                    num_ctx=8192, timeout=120)
    sql = sql.replace("```sql", "").replace("```", "").strip()
    if sql.endswith(";"): sql = sql[:-1]
    return sql


# ══════════════════════════════════════════════════════════════
# SUMMARISE — uses Mistral (FAST_MODEL)
# ══════════════════════════════════════════════════════════════

def summarise(question, result):
    prompt = f"""You are an expert Sales Analytics Assistant for a retail company.
The user asked: '{question}'
The database returned: {result.to_string()}

Rules:
- Respond in strictly 2 sentences like a business analyst presenting findings
- Be insightful — mention what the number means for the business
- Use exact names and numbers from the data
- Never add labels like 'row 0' or 'index 0' before names
- Never make up numbers not in the data
- Do not start with 'I' or 'According to'
- The leftmost column is the INDEX not the answer
- If result shows a single number use it directly
- If result has multiple rows present all clearly one by one
- If result has two columns that look like a comparison say: first period had X, second period had Y
- If discount value is between 0 and 1 multiply by 100 and show as percentage
- Never say 'round' as a column name always use the alias name.
- If the question asks "what products" or "which products" — list them simply, no insights needed
- If the question is a simple listing question — just list the items clearly, no recommendations needed
- If result contains age column, always mention the age for each person
- If result contains prime_status column, always mention prime or non-prime for each person
- Present ALL columns in the result, not just the first two"""

    # Use FAST_MODEL (mistral) for summarisation
    return call_model(prompt, model=FAST_MODEL, temperature=0.4,
                     num_ctx=4096, timeout=120)


# ══════════════════════════════════════════════════════════════
# RUN QUERY — SQL fixes and execution
# ══════════════════════════════════════════════════════════════

def run_query(sql, question=""):
    try:
        sql = re.sub(r'ROUND\(([^,]+?)(?<!::numeric),\s*2\)', lambda m: f'ROUND(({m.group(1)})::numeric, 2)', sql, flags=re.IGNORECASE)
        sql = re.sub(r'ROUND\((.+?\* 100\))::numeric,\s*2\)', lambda m: f'ROUND(({m.group(1)})::numeric, 2)', sql, flags=re.IGNORECASE)
        sql = re.sub(r'\bOrders\b', 'orders', sql)
        sql = re.sub(r'\bReturns\b', 'returns', sql)
        sql = re.sub(r'\bCustomers\b', 'customers', sql)
        sql = re.sub(r"INTERVAL '(\d+) (month|year|day|week)'", lambda m: f"INTERVAL '{m.group(1)} {m.group(2)}s'", sql, flags=re.IGNORECASE)

        if re.search(r'EXTRACT\s*\(\s*MONTH\s+FROM', sql, re.IGNORECASE):
            if not re.search(r'EXTRACT\s*\(\s*YEAR\s+FROM', sql, re.IGNORECASE):
                if not re.search(r'DATE_TRUNC', sql, re.IGNORECASE):
                    current_year = datetime.now().year
                    sql = re.sub(r'\bWHERE\b', f'WHERE EXTRACT(YEAR FROM order_date) = {current_year} AND', sql, count=1, flags=re.IGNORECASE)

        def fix_extract_month(m):
            col = m.group(1)
            month_num = int(m.group(2))
            current_month = datetime.now().month
            months_ago = (current_month - month_num) % 12
            if months_ago == 0: months_ago = 12
            return (f"{col} >= DATE_TRUNC('month', CURRENT_DATE - INTERVAL '{months_ago} months') "
                   f"AND {col} < DATE_TRUNC('month', CURRENT_DATE - INTERVAL '{months_ago - 1} months')")

        sql = re.sub(r"EXTRACT\s*\(\s*MONTH\s+FROM\s+(\w+)\s*\)\s*=\s*(\d+)", fix_extract_month, sql, flags=re.IGNORECASE)
        sql = re.sub(r"DATE_TRUNC\(('(?:quarter|month|year|week|day)'),\s*'(\d{4}-\d{2}-\d{2})'\)", lambda m: f"DATE_TRUNC({m.group(1)}, '{m.group(2)}'::date)", sql, flags=re.IGNORECASE)

        if any(word in question.lower() for word in ["cumulative", "running total", "accumulated"]):
            sql = re.sub(r'\bSUM\((\w+)\)\s+OVER\s*\(', r'SUM(SUM(\1)) OVER (', sql, flags=re.IGNORECASE)
            print(f"  🔧 Fixed: Corrected cumulative window function")

        if re.search(r'GROUP BY\s+\d+\s*,\s*\d+', sql, re.IGNORECASE):
            if re.search(r'SUM\s*\(', sql, re.IGNORECASE):
                sql = re.sub(r'GROUP BY[\s\d,]+', '', sql, flags=re.IGNORECASE)

        if re.search(r'\bJOIN\b', sql, re.IGNORECASE) and table_columns:
            group_by_match = re.search(r'GROUP BY\s+(\w+)', sql, re.IGNORECASE)
            if group_by_match:
                col = group_by_match.group(1).lower()
                tables_with_col = [t for t, cols in table_columns.items() if col in cols]
                if len(tables_with_col) > 1:
                    main_table_match = re.search(r'FROM\s+(\w+)\s+(\w+)', sql, re.IGNORECASE)
                    if main_table_match:
                        main_alias = main_table_match.group(2)
                        sql = re.sub(rf'\bSELECT\s+{col}\b', f'SELECT {main_alias}.{col}', sql, flags=re.IGNORECASE)
                        sql = re.sub(rf'\bGROUP BY\s+{col}\b', f'GROUP BY {main_alias}.{col}', sql, flags=re.IGNORECASE)

        if table_columns:
            sql_keywords = {'where','from','join','on','and','or','not','in','is','as','by','order','group','having','select','distinct','count','sum','avg','max','min','round','extract','case','when','then','else','end','null','true','false','limit','offset','union','all','inner','outer','left','right','full','cross','natural','with'}
            for alias_match in re.finditer(r'\b([a-zA-Z_]\w*)\.(\w+)', sql):
                alias = alias_match.group(1)
                col = alias_match.group(2).lower()
                if alias.lower() in sql_keywords: continue
                from_match = re.search(rf'(?:FROM|JOIN)\s+(\w+)\s+{re.escape(alias)}\b', sql, re.IGNORECASE)
                if from_match:
                    table = from_match.group(1).lower()
                    if table in table_columns and col not in table_columns[table]:
                        for other_table, cols in table_columns.items():
                            if col in cols and other_table != table:
                                other_alias_match = re.search(rf'(?:FROM|JOIN)\s+{other_table}\s+([a-zA-Z_]\w*)\b', sql, re.IGNORECASE)
                                if other_alias_match:
                                    correct_alias = other_alias_match.group(1)
                                    sql = sql.replace(f'{alias}.{col}', f'{correct_alias}.{col}')
                                    break

        for normalized, quoted in special_columns.items():
            pattern = r'(?<!")(?<!\w)' + re.escape(normalized) + r'(?!\w)(?!")'
            sql = re.sub(pattern, quoted, sql, flags=re.IGNORECASE)

        sql = re.sub(r'ROUND\(SUM\((\w+)\)::numeric,\s*2\)(?!\s+AS)', lambda m: f'ROUND(SUM({m.group(1)})::numeric, 2) AS total_{m.group(1)}', sql, flags=re.IGNORECASE)

        breakdown_words = [" by ", "by category", "by region", "by segment", "by state", "by city", "by year", "by month", "by quarter", "by ship mode", "trends", "monthly", "quarterly"]
        limit_words = ["highest", "lowest", "most", "least", "top", "bottom", "best", "worst", "maximum", "minimum"]
        is_breakdown = any(word in question.lower() for word in breakdown_words)
        is_specific = any(word in question.lower() for word in limit_words)
        if is_breakdown and not is_specific:
            sql = re.sub(r'LIMIT\s+1', '', sql, flags=re.IGNORECASE)

        group_match = re.search(r'GROUP BY (\w+)', sql, re.IGNORECASE)
        select_match = re.search(r'SELECT (.+?) FROM', sql, re.IGNORECASE)
        if group_match and select_match:
            group_col = group_match.group(1)
            select_cols = select_match.group(1)
            if group_col.lower() not in select_cols.lower():
                sql = sql.replace('SELECT ', f'SELECT {group_col}, ', 1)

        with engine.connect() as conn:
            result = pd.read_sql(text(sql), conn)
            for col in result.columns:
                if 'discount' in col.lower():
                    if result[col].between(0, 1).all():
                        result[col] = (result[col] * 100).round(2)
                        result = result.rename(columns={col: col + '_%'})
            return result

    except Exception as e:
        print(f"  ❌ SQL Error: {e}")
        return None


# ══════════════════════════════════════════════════════════════
# VISUALISATION
# ══════════════════════════════════════════════════════════════

def needs_visualisation(question):
    no_chart_phrases = [
        "give me details", "show me details", "tell me about",
        "what is this database", "how can you help", "what are you",
        "what do you do", "how many unique", "what is the total",
        "who is the", "what is the name", "what is this",
        "which customer", "which product", "which state",
        "which city", "which region", "which category"
    ]
    question_lower = question.lower()
    if any(phrase in question_lower for phrase in no_chart_phrases):
        return False

    always_chart_phrases = [
        "by region", "by category", "by segment", "by state", "by city",
        "by month", "by quarter", "by year", "trend", "compare", "comparison",
        "breakdown", "distribution", "percentage", "percent", "chart", "graph",
        "plot", "visualise", "visualize", "monthly", "quarterly", "yearly",
        "growth", "pattern", "refund", "sales by", "profit by", "orders by",
        "return rate by", "return reasons", "top ", "bottom ", "show chart",
        "with chart", "vs", "versus", "cumulative", "running total",
        "by age", "by sex", "by channel", "by device", "by payment"
    ]
    if any(phrase in question_lower for phrase in always_chart_phrases):
        return True

    return False


def create_visualisation(result, question):
    global latest_chart_buf
    try:
        if len(result.columns) < 2:
            return None

        if len(result) == 1 and len(result.columns) == 2:
            col1, col2 = result.columns[0], result.columns[1]
            val1 = pd.to_numeric(result[col1].iloc[0], errors='coerce')
            val2 = pd.to_numeric(result[col2].iloc[0], errors='coerce')
            if pd.notna(val1) and pd.notna(val2):
                result = pd.DataFrame({
                    'period': [col1.replace('_', ' ').title(), col2.replace('_', ' ').title()],
                    'value': [val1, val2]
                })

        if len(result) < 2:
            return None

        fig, ax = plt.subplots(figsize=(12, 6))
        fig.patch.set_facecolor('#1c1915')
        ax.set_facecolor('#1c1915')

        x_col = result.columns[0]
        y_col = result.columns[1]
        x_values = result[x_col].astype(str)
        y_values = pd.to_numeric(result[y_col], errors='coerce')
        question_lower = question.lower()

        def format_x_label(val):
            try:
                if re.match(r'^\d{4}-\d{2}$', val):
                    dt = datetime.strptime(val, '%Y-%m')
                    return dt.strftime('%b %Y')
                return val
            except Exception:
                return val

        x_values = x_values.apply(format_x_label)

        # PIE CHART
        if any(word in question_lower for word in ["pie", "percentage", "percent",  "contribution"]):
            numeric_col = None
            for col in result.columns:
                if pd.to_numeric(result[col], errors='coerce').notna().all():
                    numeric_col = col
                    break
            if numeric_col is None:
                return None
            if len(result) == 1:
                pct = min(float(result[numeric_col].iloc[0]), 100)
                values = [pct, 100 - pct]
                labels = [f'Yes ({pct:.1f}%)', f'No ({100-pct:.1f}%)']
                colors = ['#E85D04', '#1D9E75']
            else:
                values = pd.to_numeric(result[numeric_col], errors='coerce')
                labels = result[result.columns[0]].astype(str)
                colors = ['#2E75B6', '#1D9E75', '#E85D04', '#534AB7', '#BA7517',
                         '#C62828', '#00838F', '#558B2F', '#6A1B9A', '#2E7D32']
            wedges, texts, autotexts = ax.pie(
                values, labels=labels, colors=colors[:len(values)],
                autopct='%1.1f%%', startangle=90,
                wedgeprops={'edgecolor': 'white', 'linewidth': 2}
            )
            for text in texts: text.set_color('#e2e8f0')
            for autotext in autotexts: autotext.set_color('#ffffff')
            ax.set_title(question.title(), fontsize=14, fontweight='bold', pad=20, color='#e2e8f0')
            plt.tight_layout()
            buf = io.BytesIO()
            plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
            plt.close()
            buf.seek(0)
            latest_chart_buf = buf
            print(f"\n  📊 Pie chart generated!")
            return True

        # AREA CHART — before line chart
        if any(word in question_lower for word in ["cumulative", "running total", "accumulated"]):
            ax.fill_between(range(len(x_values)), y_values, alpha=0.4, color='#2E75B6')
            ax.plot(range(len(x_values)), y_values, color='#2E75B6', linewidth=2)
            ax.set_xticks(range(len(x_values)))
            ax.set_xticklabels(x_values, rotation=45, ha='right', color='#aaaaaa')
            chart_type = "Area"

        # LINE CHART
        elif any(word in question_lower for word in ["trend", "monthly", "quarterly", "yearly", "over time", "by month", "by year", "by quarter"]):
            ax.plot(x_values, y_values, marker='o', linewidth=2.5,
                   color='#2E75B6', markersize=8, markerfacecolor='white', markeredgewidth=2)
            ax.fill_between(range(len(x_values)), y_values, alpha=0.1, color='#2E75B6')
            chart_type = "Line"

        # SCATTER CHART
        elif any(word in question_lower for word in ["correlation", "relationship", "scatter"]):
            ax.scatter(x_values, y_values, color='#2E75B6', alpha=0.6, s=100, edgecolors='white')
            chart_type = "Scatter"

        # GROUPED BAR
        elif any(word in question_lower for word in ["vs", "versus", "compare", "comparison"]) and len(result.columns) >= 3:
            x = range(len(result))
            width = 0.35
            col_colors = ['#2E75B6', '#E85D04', '#1D9E75', '#534AB7']
            for i, col in enumerate(result.columns[1:3]):
                vals = pd.to_numeric(result[col], errors='coerce')
                offset = (i - 0.5) * width
                bars = ax.bar([xi + offset for xi in x], vals, width,
                             label=col.replace('_', ' ').title(),
                             color=col_colors[i], edgecolor='white')
                for bar, val in zip(bars, vals):
                    ax.text(bar.get_x() + bar.get_width()/2., bar.get_height(),
                           f'{val:,.0f}', ha='center', va='bottom', fontsize=8, fontweight='bold', color='#ffffff')
            ax.set_xticks(list(x))
            ax.set_xticklabels(result.iloc[:, 0].astype(str), rotation=45, ha='right', color='#aaaaaa')
            ax.legend(labelcolor='#aaaaaa', facecolor='#1c1915', edgecolor='rgba(255,255,255,0.1)')
            chart_type = "Grouped Bar"

        # HORIZONTAL BAR
        elif any(word in question_lower for word in ["top", "bottom"]) and len(result) > 3:
            sorted_df = result.copy()
            sorted_df[y_col] = pd.to_numeric(sorted_df[y_col], errors='coerce')
            sorted_df = sorted_df.sort_values(by=y_col, ascending=True)
            x_values = sorted_df[x_col].astype(str).apply(format_x_label)
            y_values = sorted_df[y_col]
            colors = plt.cm.Blues([0.3 + 0.7 * i / len(sorted_df) for i in range(len(sorted_df))])
            bars = ax.barh(x_values, y_values, color=colors, edgecolor='white')
            for bar, val in zip(bars, y_values):
                ax.text(bar.get_width(), bar.get_y() + bar.get_height()/2.,
                       f' {val:,.2f}', ha='left', va='center', fontsize=9, fontweight='bold', color='#ffffff')
            ax.tick_params(axis='y', colors='#aaaaaa')
            chart_type = "Horizontal Bar"

        # VERTICAL BAR (default)
        else:
            colors = ['#2E75B6', '#1D9E75', '#E85D04', '#534AB7', '#BA7517',
                     '#C62828', '#00838F', '#558B2F', '#6A1B9A', '#2E7D32']
            bars = ax.bar(x_values, y_values, color=colors[:len(result)],
                         edgecolor='white', linewidth=0.5, width=0.6)
            for bar, val in zip(bars, y_values):
                ax.text(bar.get_x() + bar.get_width()/2., bar.get_height(),
                       f'{val:,.2f}', ha='center', va='bottom', fontsize=9, fontweight='bold', color='#ffffff')
            chart_type = "Bar"

        ax.set_title(question.title(), fontsize=14, fontweight='bold', pad=20, color='#e2e8f0')

        if chart_type == "Horizontal Bar":
            ax.set_xlabel(y_col.replace('_', ' ').title(), fontsize=11, color='#aaaaaa')
            ax.set_ylabel(x_col.replace('_', ' ').title(), fontsize=11, color='#aaaaaa')
            ax.grid(axis='x', alpha=0.15, linestyle='--', color='white')
        else:
            ax.set_xlabel(x_col.replace('_', ' ').title(), fontsize=11, color='#aaaaaa')
            ax.set_ylabel(y_col.replace('_', ' ').title(), fontsize=11, color='#aaaaaa')
            ax.grid(axis='y', alpha=0.15, linestyle='--', color='white')

        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.spines['left'].set_color('#333333')
        ax.spines['bottom'].set_color('#333333')
        ax.tick_params(colors='#aaaaaa')
        plt.xticks(rotation=45, ha='right', color='#aaaaaa')
        plt.yticks(color='#aaaaaa')
        plt.tight_layout()

        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        plt.close()
        buf.seek(0)
        latest_chart_buf = buf
        print(f"\n  📊 {chart_type} chart generated!")
        return True

    except Exception as e:
        print(f"  ⚠️ Could not generate chart: {e}")
        return None


# ══════════════════════════════════════════════════════════════
# WORD DOC EXPORT
# ══════════════════════════════════════════════════════════════

def save_to_word(context_history):
    if not DOCX_AVAILABLE:
        return "  ⚠️  python-docx not installed. Run: pip install python-docx"
    if not context_history:
        return "  ⚠️  No conversation to save yet. Ask some questions first!"

    try:
        doc = Document()
        title = doc.add_heading('Superstore Sales Chat Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para = doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        total_para = doc.add_paragraph(f'Total Questions: {len(context_history)}')
        total_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('')

        doc.add_heading('Chat Summary', level=1)
        summary_table = doc.add_table(rows=1, cols=2)
        summary_table.style = 'Light Grid Accent 1'
        hdr = summary_table.rows[0].cells
        hdr[0].text = 'Question'
        hdr[1].text = 'Answer'
        for cell in hdr:
            cell.paragraphs[0].runs[0].bold = True

        for exchange in context_history:
            row = summary_table.add_row().cells
            row[0].text = exchange['question']
            row[1].text = exchange['answer'].split('\n\n  💡')[0].strip()

        doc.add_paragraph('')
        doc.add_heading('Detailed Analysis', level=1)

        for i, exchange in enumerate(context_history):
            doc.add_heading(f'Q{i+1}: {exchange["question"]}', level=2)
            doc.add_paragraph(exchange['answer'].split('\n\n  💡')[0].strip())
            sql_para = doc.add_paragraph()
            sql_run = sql_para.add_run(f'SQL: {exchange.get("sql", "")}')
            sql_run.font.size = Pt(8)
            sql_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            doc.add_paragraph('─' * 60)

        filename = f"chat_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(os.path.dirname(os.path.abspath(__file__)), filename)
        doc.save(filepath)
        return f"  📄 Report saved: {filename} ({len(context_history)} questions)"

    except Exception as e:
        return f"  ❌ Could not save Word doc: {e}"


# ══════════════════════════════════════════════════════════════
# MAIN ask() FUNCTION — Clean intent-based routing
# ══════════════════════════════════════════════════════════════

def ask(question):
    global context
    print(f"\n  ┌─ Question: {question}")

    # ── Database info ─────────────────────────────────────────
    database_info_phrases = [
        "what is this database", "what database", "about this database",
        "what data", "what does this contain", "what can you answer",
        "tell me about this database", "what do you know",
        "what is the database name", "database name",
        "what are you", "what do you do", "how can you help",
        "what is this", "which database"
    ]
    if any(phrase in question.lower() for phrase in database_info_phrases):
        return """This is the Superstore Sales Database (2022-2026):

  Orders     — 9,994 sales transactions
  Returns    — 1,999 product returns
  Customers  — 793 customers with demographics (age, sex, prime status, payment, channel, device)
  Geography  — 4 regions, 49 US states
  Products   — Furniture, Technology, Office Supplies

Ask me anything about sales, profit, returns, customer demographics, regions, products and trends!"""

    # ── Export trigger ────────────────────────────────────────
    export_phrases = ["download", "save", "export", "word doc", "word document", "generate report"]
    real_ctx = [ex for ex in context if ex.get("question") != "__summary__"]

    if any(phrase in question.lower() for phrase in export_phrases):
        return save_to_word(real_ctx) if real_ctx else "No conversation to save yet!"

    # ── SQL override — only when no context exists ────────────
    is_sql_question = False
    if not real_ctx:
        sql_override_phrases = [
        "what products", "which products", "list products",
        "what did", "what has", "what have", "which items",
        "how many", "how much", "what are", "show me",
        "what was", "who bought", "has bought", "have bought",
        "has purchased", "have purchased", "has ordered",
        "purchased by", "bought by", "ordered by",
        "give me all details", "all details", "details of",
        "tell me about", "info on", "information on"
        ]
        is_sql_question = any(phrase in question.lower() for phrase in sql_override_phrases)

    # ── Customer recommendation trigger ──────────────────────
    # Only fires if NOT a SQL/data question
    recommendation_triggers = [
        "recommend for", "recommendation for",
        "you may also", "you may buy",
        "suggest products for",
        "what would you recommend for",
        "what should we recommend for",
        "other customers also bought",
        "similar customers bought"
    ]

    if not is_sql_question and any(trigger in question.lower() for trigger in recommendation_triggers):
        try:
            with engine.connect() as conn:
                customers = conn.execute(text(
                    "SELECT customer_name FROM customers"
                )).fetchall()
                customer_names = [row[0] for row in customers]

            found_customer = None
            q_lower = question.lower()
            for name in customer_names:
                if name.lower() in q_lower:
                    found_customer = name
                    break

            if found_customer:
                print(f"  🛍️ Customer recommendation for: {found_customer}")
                return handle_customer_recommendation(question, found_customer)
        except Exception as e:
            print(f"  ⚠️ Customer lookup failed: {e}")

    # ── Classify intent ───────────────────────────────────────
    intent = classify_intent(question)
    print(f"  🎯 Intent: {intent}")

    # ── Route to handler ──────────────────────────────────────
    if intent == "generic":
        return handle_generic(question)

    elif intent == "contextual":
        # Pass full context (includes summary entry) — handle_contextual extracts both
        return handle_contextual(question, context)

    elif intent == "recommendation":
        answer, relevant_data = handle_recommendation(question, context)
        context.append({"question": question, "sql": "", "answer": answer, "result": relevant_data})
        if len(context) > 10: context.pop(0)
        return answer

    elif intent == "subjective_sql":
        return handle_subjective_sql(question, context)

    else:  # sql (default)
        return handle_sql(question)


# ══════════════════════════════════════════════════════════════
# TERMINAL LOOP
# ══════════════════════════════════════════════════════════════

if __name__ == "__main__":
    print()
    print("  Welcome to the Superstore Sales Database Chatbot!")
    print()
    print("  Commands: 'exit' | 'clear' | 'context' | 'download'")
    print()
    print("=" * 60)

    while True:
        print()
        question = input("  You: ")

        if question.strip().lower() == "exit":
            print("  Goodbye!")
            break
        if question.strip().lower() == "clear":
            context = []
            print("  ✅ Conversation history cleared!")
            continue
        if question.strip().lower() == "context":
            if not context:
                print("  No conversation history yet!")
            else:
                for i, ex in enumerate(context):
                    print(f"\n  {i+1}. Q: {ex['question']}\n     A: {ex['answer']}")
            continue
        if question.strip() == "":
            print("  ⚠️  Please type a question!")
            continue

        print("  🤖 Thinking...")
        answer = ask(question)
        print(f"\n  └─ Answer:")
        wrapped = textwrap.fill(answer, width=70, initial_indent="     ", subsequent_indent="     ")
        print(wrapped)
        print()
        print("  " + "─" * 56)